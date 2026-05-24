"""Gera todos os documentos entregáveis da challenge Ford×FIAP 2026:
  1. Relatório PDF do Desafio 2 (ML)
  2. Business Canvas (.docx)
  3. Quadro de Valor (.docx)
  4. README dos entregáveis (.docx)

Saída: docs/deliverables/
"""
from pathlib import Path
import json

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak,
)

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "deliverables"
OUT.mkdir(parents=True, exist_ok=True)

# ===== Time =====
EQUIPE = [
    ("Guilherme", "554962"),
    ("Pedro",     "555556"),
    ("Fabrício",  "558216"),
    ("Vitor",     "554893"),
    ("Matheus",   "555447"),
]
PROJETO = "Faro AI — Inteligência Competitiva + Retenção VIN Share"
EMPRESA = "Faro AI"
TAGLINE = "AI que tem faro pro cliente certo."
CHALLENGE = "Ford × FIAP 2026 · 1ª Sprint"
ENTREGA = "24/05/2026"

# Ford brand colors
FORD_BLUE = colors.HexColor("#003478")
FORD_BLUE_LIGHT = colors.HexColor("#0066B2")
FORD_GREY = colors.HexColor("#4A4A4A")
FORD_AMBER = colors.HexColor("#FFA500")

# Métricas: prioriza métricas REAIS Ford BR (175k VINs), com fallback ao sintético.
METRICS = json.loads((ROOT / "services" / "ml" / "models" / "metrics.json").read_text(encoding="utf-8"))
_real_path = ROOT / "services" / "ml" / "models" / "metrics_real.json"
METRICS_REAL = json.loads(_real_path.read_text(encoding="utf-8")) if _real_path.exists() else None

# D1 schema canônico (262 itens × 14 seções)
_d1_path = ROOT / "services" / "ml" / "data" / "ford-d1-ranger-26my.json"
D1_SCHEMA = json.loads(_d1_path.read_text(encoding="utf-8")) if _d1_path.exists() else None


# ============================================================
# 1. RELATÓRIO D2 (PDF) — ML segmentation + classification
# ============================================================

def build_ml_report():
    pdf_path = OUT / "Relatorio_Desafio_2_ML.pdf"
    doc = SimpleDocTemplate(
        str(pdf_path), pagesize=A4,
        leftMargin=2.2 * cm, rightMargin=2.2 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    H1 = ParagraphStyle("H1", parent=styles["Heading1"], fontName="Helvetica-Bold",
                        fontSize=22, textColor=FORD_BLUE, spaceAfter=14, leading=26)
    H2 = ParagraphStyle("H2", parent=styles["Heading2"], fontName="Helvetica-Bold",
                        fontSize=15, textColor=FORD_BLUE, spaceBefore=18, spaceAfter=8, leading=18)
    H3 = ParagraphStyle("H3", parent=styles["Heading3"], fontName="Helvetica-Bold",
                        fontSize=12, textColor=FORD_BLUE_LIGHT, spaceBefore=10, spaceAfter=4)
    BODY = ParagraphStyle("Body", parent=styles["BodyText"], fontName="Helvetica",
                          fontSize=10.5, leading=15, alignment=TA_JUSTIFY, spaceAfter=8)
    CAPTION = ParagraphStyle("Cap", parent=BODY, fontSize=9, textColor=FORD_GREY,
                             alignment=TA_CENTER, spaceAfter=14)
    META = ParagraphStyle("Meta", parent=BODY, fontSize=10, alignment=TA_LEFT, spaceAfter=4)

    # Estilos para conteúdo de tabela (com word-wrap)
    TBL_HEAD = ParagraphStyle("TblHead", parent=BODY, fontName="Helvetica-Bold",
                              fontSize=10, textColor=colors.white, alignment=TA_LEFT, leading=12, spaceAfter=0)
    TBL_HEAD_C = ParagraphStyle("TblHeadC", parent=TBL_HEAD, alignment=TA_CENTER)
    TBL_CELL = ParagraphStyle("TblCell", parent=BODY, fontName="Helvetica",
                              fontSize=9.5, alignment=TA_LEFT, leading=12, spaceAfter=0)
    TBL_CELL_C = ParagraphStyle("TblCellC", parent=TBL_CELL, alignment=TA_CENTER)
    TBL_CELL_MONO = ParagraphStyle("TblMono", parent=TBL_CELL, fontName="Courier",
                                   fontSize=8.5, leading=11)

    def P(txt, style=TBL_CELL):
        """Helper: paragraph wrapped cell."""
        return Paragraph(str(txt), style)
    def PH(txt, center=False):
        return Paragraph(str(txt), TBL_HEAD_C if center else TBL_HEAD)

    story = []

    # ---- Capa ----
    story.append(Spacer(1, 4 * cm))
    story.append(Paragraph(PROJETO, H1))
    story.append(Paragraph("Relatório Técnico — Desafio 2", H2))
    story.append(Paragraph("Segmentação Comportamental e Classificação Preditiva de Clientes Ford", BODY))
    story.append(Spacer(1, 2 * cm))
    story.append(Paragraph(f"<b>Challenge:</b> {CHALLENGE}", META))
    story.append(Paragraph(f"<b>Data de entrega:</b> {ENTREGA}", META))
    story.append(Paragraph("<b>Equipe (FIAP):</b>", META))
    for nome, rm in EQUIPE:
        story.append(Paragraph(f"&nbsp;&nbsp;• {nome} — RM {rm}", META))
    story.append(PageBreak())

    # ---- Sumário Executivo ----
    story.append(Paragraph("1. Sumário Executivo", H2))
    story.append(Paragraph(
        "Este relatório apresenta a solução analítica desenvolvida para o <b>Desafio 2</b> "
        "do Ford×FIAP Challenge 2026: aumentar a retenção de clientes na rede oficial de "
        "manutenção pós-venda da Ford. O projeto combina técnicas de aprendizado não-"
        "supervisionado (clustering) sobre o histórico completo de clientes (Base 1) com "
        "um classificador supervisionado (XGBoost) treinado exclusivamente sobre dados "
        "disponíveis no momento da compra (Base 2), garantindo zero data leakage.", BODY))
    story.append(Paragraph(
        "A solução identifica <b>quatro perfis comportamentais</b> distintos — Cliente Fiel, "
        "Cliente Esquecido, Cliente Econômico e Cliente de Abandono — e permite, no ato da "
        "venda, prever a qual perfil o novo cliente provavelmente pertencerá. Para cada perfil "
        "são propostas ações de retenção específicas, executáveis pelo CRM da concessionária.", BODY))

    if METRICS_REAL:
        n_train = METRICS_REAL.get("n_samples_train", 0)
        n_test = METRICS_REAL.get("n_samples_test", 0)
        total = n_train + n_test
        story.append(Paragraph(
            f"<b>Upgrade: dados reais Ford BR.</b> O modelo de produção (<code>{METRICS_REAL['model_version']}</code>) "
            f"foi treinado com <b>{total:,}</b> VINs reais do arquivo <code>vin_share_Desafio_02.xlsx</code> "
            f"({n_train:,} treino · {n_test:,} teste), entregando "
            f"<b>accuracy {METRICS_REAL['accuracy']:.1%}</b> · F1-weighted {METRICS_REAL['f1_weighted']:.3f} · "
            f"F1-macro {METRICS_REAL['f1_macro']:.3f}. As métricas no sintético (Seção 5) são mantidas como "
            "baseline histórico de desenvolvimento — todo o produto consulta o modelo real.".replace(',', '.'), BODY))

    # ---- Hipótese de negócio ----
    story.append(Paragraph("2. Hipótese de Negócio", H2))
    story.append(Paragraph(
        "A área de pós-venda da Ford parte da hipótese de que existem quatro arquétipos "
        "principais entre os compradores de veículos novos:", BODY))

    perfis_hip = [
        [PH("Perfil"),    PH("Comportamento esperado")],
        [P("Fiel"),       P("Retorna consistentemente à rede oficial, independente de preço.")],
        [P("Abandono"),   P("Faz no máximo a primeira revisão e migra para oficinas externas.")],
        [P("Esquecido"),  P("Perde o timing da revisão; tenta voltar tarde e se frustra.")],
        [P("Econômico"),  P("Mantém vínculo mas é altamente sensível a preço/promoção.")],
    ]
    t = Table(perfis_hip, colWidths=[3.5 * cm, 13.0 * cm], repeatRows=1)
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), FORD_BLUE),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F0F4F8')]),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#CCCCCC')),
    ]))
    story.append(t)
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph(
        "Coube ao grupo <b>validar</b> se os dados sustentam essa estrutura via clustering "
        "não-supervisionado, sem assumir as classes a priori.", BODY))

    # ---- Bases de dados ----
    story.append(Paragraph("3. Bases de Dados", H2))
    story.append(Paragraph(
        "Durante o desenvolvimento inicial, usamos bases sintéticas geradas em "
        "<code>services/ml/src/synthetic.py</code> para validar todo o pipeline (EDA, "
        "clustering, treino, avaliação) sem depender da chegada do dataset oficial. "
        "<b>Na versão final em produção</b>, o modelo <code>xgb-real-v1</code> foi "
        "treinado com a <b>base real Ford BR de 175.554 VINs</b> "
        "(<code>vin_share_Desafio_02.xlsx</code>) — descrito em detalhe na Seção 5.bis. "
        "As métricas sintéticas (Seção 5) são mantidas apenas como baseline histórico de "
        "desenvolvimento. As bases sintéticas seguiram a mesma estrutura conceitual "
        "Base 1 (histórico completo) e Base 2 (apenas dados pré-compra):", BODY))

    bases_t = [
        [PH("Base"), PH("Conteúdo"), PH("Uso")],
        [P("Base 1"),
         P("Histórico completo (incluindo comportamento pós-compra: revisões realizadas, gasto, dias até última visita, churn)"),
         P("Exclusivamente para clustering — define ground truth dos perfis")],
        [P("Base 2"),
         P("Apenas variáveis disponíveis no ato da compra (demográficas, score, modelo, forma de pagamento, canal de aquisição, primeiro carro, test drive)"),
         P("Treino do classificador supervisionado")],
    ]
    t2 = Table(bases_t, colWidths=[2.0 * cm, 8.5 * cm, 6.0 * cm], repeatRows=1)
    t2.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), FORD_BLUE),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F0F4F8')]),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#CCCCCC')),
    ]))
    story.append(t2)
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph(
        "<b>Regra crítica observada:</b> nenhuma variável que represente comportamento "
        "pós-compra (revisões realizadas, gasto acumulado, churn) entrou no classificador. "
        "A separação Base1/Base2 é implementada em <code>synthetic.split_base1_base2()</code> "
        "e a lista de features do classificador é fixa em <code>classifier.py</code>:", BODY))
    story.append(Paragraph(
        "<i>idade, genero, regiao, renda_mensal_brl, estado_civil, score_credito, "
        "modelo_comprado, versao_comprada, preco_pago_brl, financiamento, parcelas, "
        "canal_aquisicao, primeiro_carro, test_drive_realizado.</i>", BODY))

    # ---- Pipeline ----
    story.append(Paragraph("4. Pipeline Analítico", H2))
    story.append(Paragraph("4.1 Preparação", H3))
    story.append(Paragraph(
        "EDA executada no notebook <code>services/ml/notebooks/ford_segmentation.ipynb</code> "
        "cobre: distribuições das variáveis, correlações, valores faltantes (zero, dado "
        "sintético controlado), boxplots por perfil-real e validação de cardinalidade das "
        "categóricas. As features numéricas são padronizadas (<code>StandardScaler</code>) "
        "e as categóricas são one-hot (<code>OneHotEncoder</code> com <code>handle_unknown='ignore'</code>) "
        "via <code>ColumnTransformer</code>.", BODY))

    story.append(Paragraph("4.2 Segmentação (Base 1)", H3))
    story.append(Paragraph(
        "Algoritmo: <b>K-Means</b> com <code>k=4</code> escolhido por elbow method "
        "(SSE) e validado por <i>silhouette score</i>. Variáveis usadas: idade, renda, "
        "score, preço pago, parcelas, financiamento, canal de aquisição "
        "<i>combinadas com</i> variáveis comportamentais (revisões realizadas, dias até "
        "última manutenção, gasto acumulado, taxa de aderência ao plano).", BODY))

    sil = METRICS["cluster_metrics"]["silhouette_score"]
    sizes = METRICS["cluster_metrics"]["cluster_sizes"]
    mapping = METRICS["cluster_metrics"]["cluster_to_persona"]
    total = sum(sizes.values())

    cluster_t = [[PH("Cluster", center=True), PH("Persona", center=True),
                  PH("N clientes", center=True), PH("% da base", center=True)]]
    for cid, persona in mapping.items():
        n = sizes.get(cid, 0)
        cluster_t.append([P(cid, TBL_CELL_C), P(persona.capitalize(), TBL_CELL_C),
                          P(str(n), TBL_CELL_C), P(f"{100*n/total:.1f}%", TBL_CELL_C)])
    t3 = Table(cluster_t, colWidths=[2.5 * cm, 4.0 * cm, 3.0 * cm, 3.0 * cm], repeatRows=1)
    t3.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), FORD_BLUE),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F0F4F8')]),
        ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#CCCCCC')),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(t3)
    story.append(Paragraph(
        f"Silhouette score: <b>{sil:.3f}</b> (estrutura presente, separação razoável "
        "para dados socio-comportamentais reais).", CAPTION))

    story.append(Paragraph("4.3 Mapeamento Cluster → Persona", H3))
    story.append(Paragraph(
        "Implementado em <code>classifier._map_clusters_to_personas()</code>, traduz os "
        "clusters numéricos em rótulos de negócio por análise das médias intra-cluster "
        "(score de aderência, recência de visita, frequência de revisão, sensibilidade "
        "a desconto). Evita-se nomenclaturas genéricas como “Cluster 0”.", BODY))

    story.append(Paragraph("4.4 Classificação (Base 2)", H3))
    story.append(Paragraph(
        "Algoritmo: <b>XGBoost</b> (multi:softprob, 300 estimadores, max_depth=6, "
        "learning_rate=0.08, eval_metric=mlogloss). Target: o cluster atribuído a cada "
        "cliente na Base 1, propagado para o registro pareado na Base 2. "
        "Split 80/20 estratificado, <code>random_state=42</code>.", BODY))

    # ---- Resultados ----
    story.append(PageBreak())
    story.append(Paragraph("5. Resultados do Classificador", H2))

    clsm = METRICS["classifier_metrics"]
    big = [
        [PH("Métrica"), PH("Valor", center=True)],
        [P("Accuracy"),    P(f"{clsm['accuracy']:.3f}", TBL_CELL_C)],
        [P("F1 macro"),    P(f"{clsm['f1_macro']:.3f}", TBL_CELL_C)],
        [P("F1 weighted"), P(f"{clsm['f1_weighted']:.3f}", TBL_CELL_C)],
        [P("Modelo"),      P(clsm["model_version"], TBL_CELL_C)],
    ]
    t4 = Table(big, colWidths=[5 * cm, 4 * cm], repeatRows=1)
    t4.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), FORD_BLUE),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F0F4F8')]),
        ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#CCCCCC')),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(t4)

    story.append(Paragraph("5.1 Matriz de Confusão", H3))
    cm_labels = clsm["labels"]
    matrix = clsm["confusion_matrix"]
    TBL_CELL_RH = ParagraphStyle("TblCellRH", parent=TBL_CELL, fontName="Helvetica-Bold",
                                 textColor=colors.white)
    cm_rows = [[Paragraph("", TBL_HEAD_C)] +
               [PH("Pred " + l, center=True) for l in cm_labels]]
    for i, label in enumerate(cm_labels):
        cm_rows.append([Paragraph("Real " + label, TBL_CELL_RH)] +
                       [P(str(x), TBL_CELL_C) for x in matrix[i]])
    t5 = Table(cm_rows, colWidths=[3.2 * cm] + [2.6 * cm] * 4)
    t5.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), FORD_BLUE),
        ('BACKGROUND', (0, 1), (0, -1), FORD_BLUE_LIGHT),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#CCCCCC')),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))
    # diagonal verde claro
    for i in range(len(cm_labels)):
        t5.setStyle(TableStyle([
            ('BACKGROUND', (i + 1, i + 1), (i + 1, i + 1), colors.HexColor('#D4F0D4'))
        ]))
    story.append(t5)
    story.append(Spacer(1, 0.3 * cm))

    # ---- Leitura executiva ----
    story.append(Paragraph("5.2 Leitura por classe", H3))
    diag = []
    for i, label in enumerate(cm_labels):
        row_sum = sum(matrix[i])
        col_sum = sum(matrix[r][i] for r in range(len(cm_labels)))
        tp = matrix[i][i]
        recall = tp / row_sum if row_sum else 0
        precision = tp / col_sum if col_sum else 0
        diag.append((label, precision, recall, row_sum))

    diag_t = [[PH("Perfil"), PH("Precisão", center=True),
               PH("Recall", center=True), PH("Suporte", center=True)]]
    for label, pp, rr, ss in diag:
        diag_t.append([P(label.capitalize()),
                       P(f"{pp:.2f}", TBL_CELL_C),
                       P(f"{rr:.2f}", TBL_CELL_C),
                       P(str(ss), TBL_CELL_C)])
    t6 = Table(diag_t, colWidths=[4 * cm, 3 * cm, 3 * cm, 3 * cm], repeatRows=1)
    t6.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), FORD_BLUE),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F0F4F8')]),
        ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#CCCCCC')),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(t6)

    story.append(Paragraph(
        f"<b>Análise:</b> a acurácia de {clsm['accuracy']:.1%} é coerente com a sobreposição "
        "intrínseca entre perfis comportamentais que compartilham faixa demográfica. "
        "O modelo é particularmente forte em identificar o perfil <b>Fiel</b> (alta precisão "
        "e recall) — exatamente o segmento mais valioso para o programa de fidelidade. "
        "Há confusão esperada entre <b>Esquecido</b> e <b>Econômico</b>, segmentos que "
        "podem ser tratados com ações parcialmente sobrepostas (lembretes + condições "
        "comerciais). Para o uso no chão da concessionária, a probabilidade de cada "
        "classe (e não só a top-1) é exibida ao vendedor, permitindo decisão informada.", BODY))

    # ---- 5.bis Modelo em PRODUÇÃO (dados reais Ford BR) ----
    if METRICS_REAL:
        story.append(Paragraph("5.bis Modelo em Produção — Dados Reais Ford BR", H2))
        n_train = METRICS_REAL.get("n_samples_train", 0)
        n_test = METRICS_REAL.get("n_samples_test", 0)
        total = n_train + n_test
        story.append(Paragraph(
            f"A versão de produção (<code>{METRICS_REAL['model_version']}</code>) foi retreinada "
            f"sobre <b>{total:,}</b> VINs reais do arquivo <code>vin_share_Desafio_02.xlsx</code> "
            f"fornecido pela Ford. ETL em <code>scripts/etl-d2-real.py</code> converte XLSX → Parquet "
            "gerado localmente (não commitado por tamanho/governança de dados), "
            "reconstrói Base 1 (histórico completo) e Base 2 (features pré-compra: ano_modelo, "
            "ano_venda, mês_venda, dealer_venda, modelo, % histórico de cada perfil no dealer e "
            "no modelo). Implementação em <code>services/ml/src/classifier_real.py</code>.".replace(',', '.'),
            BODY))

        real_kpi = [
            [PH("Métrica"), PH("Valor", center=True)],
            [P("Accuracy"),    P(f"{METRICS_REAL['accuracy']:.4f}", TBL_CELL_C)],
            [P("F1 macro"),    P(f"{METRICS_REAL['f1_macro']:.4f}", TBL_CELL_C)],
            [P("F1 weighted"), P(f"{METRICS_REAL['f1_weighted']:.4f}", TBL_CELL_C)],
            [P("Treino / Teste"), P(f"{n_train:,} / {n_test:,}".replace(',', '.'), TBL_CELL_C)],
            [P("Modelo"),      P(METRICS_REAL["model_version"], TBL_CELL_C)],
        ]
        t_real = Table(real_kpi, colWidths=[5 * cm, 5 * cm], repeatRows=1)
        t_real.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), FORD_BLUE),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F0F4F8')]),
            ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#CCCCCC')),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(t_real)
        story.append(Spacer(1, 0.3 * cm))

        # Matriz confusão real
        story.append(Paragraph("5.bis.1 Matriz de Confusão (dados reais)", H3))
        cm_labels_r = METRICS_REAL["labels"]
        matrix_r = METRICS_REAL["confusion_matrix"]
        TBL_CELL_RH2 = ParagraphStyle("TblCellRH2", parent=TBL_CELL, fontName="Helvetica-Bold",
                                       textColor=colors.white)
        cm_rows_r = [[Paragraph("", TBL_HEAD_C)] +
                     [PH("Pred " + l, center=True) for l in cm_labels_r]]
        for i, label in enumerate(cm_labels_r):
            cm_rows_r.append([Paragraph("Real " + label, TBL_CELL_RH2)] +
                             [P(f"{x:,}".replace(',', '.'), TBL_CELL_C) for x in matrix_r[i]])
        t_real_cm = Table(cm_rows_r, colWidths=[3.2 * cm] + [2.8 * cm] * 4)
        t_real_cm.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), FORD_BLUE),
            ('BACKGROUND', (0, 1), (0, -1), FORD_BLUE_LIGHT),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#CCCCCC')),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        for i in range(len(cm_labels_r)):
            t_real_cm.setStyle(TableStyle([
                ('BACKGROUND', (i + 1, i + 1), (i + 1, i + 1), colors.HexColor('#D4F0D4'))
            ]))
        story.append(t_real_cm)
        story.append(Spacer(1, 0.3 * cm))

        story.append(Paragraph(
            f"<b>Distribuição observada na base real:</b> ~55% Esquecido, ~19% Fiel, ~15% "
            "Econômico, ~11% Abandono. A predominância de Esquecido confirma a hipótese "
            "de que o maior ganho operacional está em reativar clientes que perdem o "
            "timing da revisão — exatamente o foco do módulo <code>/leads</code>, que "
            "ranqueia por risco de evasão.", BODY))

    # ---- 5.ter Schema canônico Ford D1 (262 atributos) ----
    if D1_SCHEMA:
        n_secoes = len(D1_SCHEMA["sections"])
        n_items = sum(len(v) for v in D1_SCHEMA["sections"].values())
        story.append(Paragraph("5.ter Cobertura do Desafio 1 — Schema Canônico Ford D1", H2))
        story.append(Paragraph(
            f"O Desafio 1 (Inteligência Competitiva) pede que o sistema permita "
            "comparações ponta-a-ponta com qualquer concorrente. A Ford forneceu um "
            f"<b>template oficial de Vehicle Data com {n_items} atributos em {n_secoes} seções</b> "
            "(Wheels, Connectivity, Ice Line Up, Air conditioning, Safety, High tech, Global "
            "Closing, Trim, SunRoof, Seats, Lights, 4X4, Others + Motorização). Implementamos "
            "esse template como schema canônico no banco:", BODY))
        d1_t = [
            [PH("Objeto"), PH("Função")],
            [P("catalog_items"), P(f"{n_items} linhas read-mostly. Cada linha = 1 atributo "
                                    "do template (secao, ordem, nome, tipo: flag/numeric/text)")],
            [P("vehicle_catalog_values"), P("Bridge (vehicle × item) com valor preenchido (X / 0 / "
                                            "numérico) + confiança + fonte")],
            [P("GET /competitive/catalog-items"), P("Retorna os 262 atributos agrupados por seção")],
            [P("POST /competitive/compare/canonico"), P("Retorna matriz 262 × N veículos para "
                                                       "renderização da tabela fixa")],
            [P("/veiculos/comparar"), P("UI exibe a tabela canônica colapsável por seção, com "
                                        "ícones ✓/✗ para flags e valores numéricos por trim")],
        ]
        t_d1 = Table(d1_t, colWidths=[5.5 * cm, 11.0 * cm], repeatRows=1)
        t_d1.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), FORD_BLUE),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F0F4F8')]),
            ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#CCCCCC')),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(t_d1)
        story.append(Spacer(1, 0.3 * cm))
        story.append(Paragraph(
            f"As três versões da Ford Ranger 26MY (XLT, Limited, Limited+) já chegam com "
            f"todos os {n_items} valores preenchidos a partir do datasheet oficial. Cadastros "
            "de concorrentes (Hilux, Amarok, SW4, Frontier…) herdam o schema vazio e podem "
            "ser preenchidos manualmente ou via IA puxando do site oficial — habilitando "
            "comparação 1:1 estrita, como pede a Ford.", BODY))

    # ---- Estratégias por perfil ----
    story.append(Paragraph("6. Estratégias de Retenção por Perfil", H2))
    story.append(Paragraph(
        "Codificadas em <code>classifier.ACOES_POR_PERFIL</code> e retornadas pelo "
        "endpoint <code>POST /predict</code> junto com o perfil predito — o CRM da "
        "concessionária pode disparar a ação automaticamente ou exibir para o vendedor:", BODY))

    estrat = [
        ("Fiel",       "Convite ao programa de fidelidade premium • Oferta de upgrade no próximo modelo • Convites para eventos da marca."),
        ("Abandono",   "Contato proativo do consultor sênior em até 7 dias • Pacote de revisão com desconto agressivo (-30%) • Cashback na 1ª manutenção fora da garantia • Pesquisa qualitativa de motivo de saída."),
        ("Esquecido",  "Campanha SMS+WhatsApp lembrando próxima revisão • Bônus por trazer o carro nos próximos 30 dias • Oferta de busca/entrega domiciliar do veículo."),
        ("Econômico",  "Pacote de revisão com preço fechado • Programa de assinatura de manutenção (mensalidade baixa) • Cross-sell de peças genuínas com desconto progressivo."),
    ]
    rows = [[PH("Perfil"), PH("Ações sugeridas")]]
    for nome, txt in estrat:
        rows.append([P(nome), P(txt)])
    t7 = Table(rows, colWidths=[2.6 * cm, 13.9 * cm], repeatRows=1)
    t7.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), FORD_BLUE),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F0F4F8')]),
        ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#CCCCCC')),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 7),
        ('RIGHTPADDING', (0, 0), (-1, -1), 7),
    ]))
    story.append(t7)

    # ---- Integração ----
    story.append(Paragraph("7. Integração Operacional", H2))
    story.append(Paragraph(
        "O modelo é servido por uma FastAPI em <code>services/ml/src/main.py</code> "
        "(porta 8001). O API gateway Node consome em <code>POST /predict</code> a cada "
        "novo cadastro de cliente no portal Ford. O retorno (<i>perfil_predito, "
        "probabilidades, risco_evasao, recomendacoes_acao</i>) é gravado em "
        "<code>ai_predictions</code> e renderizado no card do cliente.", BODY))
    story.append(Paragraph(
        "<b>Segurança/Privacidade:</b> o payload enviado ao serviço de ML não contém "
        "PII — <code>dealership_id</code> é pseudonimizado via HMAC-SHA256 antes do "
        "envio, e nome/CPF/email/telefone nunca entram no modelo. A integridade do "
        "payload é validada via assinatura <code>X-Payload-Signature</code> (também "
        "HMAC-SHA256 da request body). Detalhes em <code>docs/SECURITY.md</code>.", BODY))

    # ---- Conclusão ----
    story.append(Paragraph("8. Conclusão e Próximos Passos", H2))
    story.append(Paragraph(
        "A solução entrega uma resposta funcional e auditável para o problema de "
        "retenção: identifica perfis comportamentais reais a partir dos dados, prevê "
        "o perfil de novos clientes no ato da compra e sugere ações operacionais "
        "tangíveis para cada segmento.", BODY))
    story.append(Paragraph("<b>Próximos passos sugeridos:</b>", BODY))
    for txt in [
        "Coletar dados reais adicionais de uma loja-piloto, com tracking de conversão pós-ação, "
        "para re-treino contínuo e validação operacional do modelo já em produção "
        "(<code>xgb-real-v1</code>, 175.554 VINs, accuracy 62,7%).",
        "Adicionar tracking de conversão das ações sugeridas (A/B) para fechar o loop "
        "modelo-ação-resultado e quantificar uplift de retenção por perfil.",
        "Treinar variantes específicas por região/dealership para capturar diferenças "
        "regionais de comportamento (preço-sensibilidade no Norte/Nordeste, etc.).",
        "Implementar drift detection mensal — re-treino automático quando distribuição "
        "das features mudar significativamente.",
    ]:
        story.append(Paragraph("• " + txt, BODY))

    # ---- Entregáveis técnicos ----
    story.append(Paragraph("9. Entregáveis Técnicos", H2))
    deliv = [
        [PH("Item"), PH("Localização")],
        [P("Notebook EDA + treino"),    P("services/ml/notebooks/ford_segmentation.ipynb", TBL_CELL_MONO)],
        [P("Dados sintéticos (baseline)"), P("gerados localmente: services/ml/data/base1_full.parquet, base2_classifier.parquet", TBL_CELL_MONO)],
        [P("Dados reais Ford BR"),      P("gerados localmente: services/ml/data/ford_real_base1_full.parquet, ford_real_base2_classifier.parquet", TBL_CELL_MONO)],
        [P("ETL D2 (XLSX → Parquet)"),  P("scripts/etl-d2-real.py", TBL_CELL_MONO)],
        [P("Modelo sintético"),         P("gerado localmente: services/ml/models/classifier_base2.joblib", TBL_CELL_MONO)],
        [P("Modelo PRODUÇÃO (real)"),   P("gerado localmente: services/ml/models/classifier_real_v1.joblib", TBL_CELL_MONO)],
        [P("Métricas (sintético)"),     P("geradas localmente: services/ml/models/metrics.json", TBL_CELL_MONO)],
        [P("Métricas (real)"),          P("geradas localmente: services/ml/models/metrics_real.json", TBL_CELL_MONO)],
        [P("Classifier real"),          P("services/ml/src/classifier_real.py", TBL_CELL_MONO)],
        [P("Serviço de inferência"),    P("services/ml/src/main.py (FastAPI)", TBL_CELL_MONO)],
        [P("Hybrid ML+IA classifier"),  P("apps/api/src/modules/retention/hybrid-classifier.ts", TBL_CELL_MONO)],
        [P("D1 Schema canônico (262)"), P("services/ml/data/ford-d1-ranger-26my.json + migration 015", TBL_CELL_MONO)],
        [P("D1 Populate script"),       P("scripts/populate-catalog-canonico.mjs", TBL_CELL_MONO)],
        [P("Resumo Ford real"),         P("services/ml/data/ford-real-summary.json", TBL_CELL_MONO)],
        [P("Relatório (este doc)"),     P("docs/deliverables/Relatorio_Desafio_2_ML.pdf", TBL_CELL_MONO)],
    ]
    t8 = Table(deliv, colWidths=[5.0 * cm, 11.5 * cm], repeatRows=1)
    t8.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), FORD_BLUE),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F0F4F8')]),
        ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#CCCCCC')),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
    ]))
    story.append(t8)

    story.append(Spacer(1, 0.6 * cm))
    story.append(Paragraph(
        f"<i>Documento gerado em {ENTREGA} para a {CHALLENGE}.</i>", CAPTION))

    doc.build(story)
    print(f"OK: {pdf_path.name}")


# ============================================================
# Helpers docx
# ============================================================

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color=(0x00, 0x34, 0x78)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)
    return h


def add_para(doc, text, bold=False, size=11, align=None):
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p


# ============================================================
# 2. BUSINESS CANVAS (.docx)
# ============================================================

CANVAS = {
    "Parcerias-Chave": (
        "• Ford Motor Company Brasil — sponsor, dados, validação\n"
        "• Rede de concessionárias autorizadas Ford BR\n"
        "• Supabase — backend-as-a-service (DB + Auth + RLS)\n"
        "• Anthropic, OpenAI, Google — providers de LLM\n"
        "• FIPE.online (Parallelum) — base de preços oficial\n"
        "• 411 Vehicle Data (RapidAPI) — specs técnicos US\n"
        "• Sites oficiais das fabricantes — fonte de e-books"
    ),
    "Atividades-Chave": (
        "• Engenharia da plataforma (API, web, mobile)\n"
        "• Pipeline de coleta multi-fonte de specs\n"
        "• Treino e versionamento de modelos de ML\n"
        "• Curadoria de chaves de IA por função\n"
        "• Monitoramento de qualidade dos dados\n"
        "• Treinamento de vendedores/consultores na ferramenta\n"
        "• Suporte e auditoria de segurança contínua"
    ),
    "Recursos-Chave": (
        "• Equipe técnica (5 dev FIAP + Ford TI)\n"
        "• Modelos de ML treinados (XGBoost, clustering)\n"
        "• Acesso autenticado a Supabase, RapidAPI, LLMs\n"
        "• Base sintética validada e base real Ford\n"
        "• Repositório git versionado + CI/CD\n"
        "• Cloud infra (Vercel/Fly + Supabase managed)"
    ),
    "Proposta de Valor": (
        "Para vendedores e gestores Ford:\n\n"
        "1. INTELIGÊNCIA COMPETITIVA — catálogo confiável de "
        "veículos próprios e concorrentes com comparativo "
        "imediato, citando fonte de cada spec.\n\n"
        "2. RETENÇÃO PROATIVA — prever no ato da venda o perfil "
        "comportamental do cliente e disparar a ação certa "
        "(fidelização, lembrete, desconto), aumentando o VIN "
        "Share da rede oficial.\n\n"
        "Tudo via web e mobile, com latência baixa e auditoria."
    ),
    "Relacionamento com Clientes": (
        "• Self-service via portal web (gestores/admin)\n"
        "• App mobile para consultores em loja\n"
        "• Treinamento on-demand (vídeos curtos)\n"
        "• Suporte por canal interno Ford\n"
        "• Iteração mensal com feedback de concessionários\n"
        "• Painel de observabilidade e SLAs"
    ),
    "Canais": (
        "• Portal web (apps/web) — desktop\n"
        "• App nativo iOS/Android (apps/mobile via Expo)\n"
        "• API REST documentada via Swagger (/docs)\n"
        "• Webhook para CRM da concessionária (futuro)\n"
        "• E-mail transacional para alertas críticos"
    ),
    "Segmentos de Clientes": (
        "• Vendedores e consultores das concessionárias Ford\n"
        "• Gerentes de pós-venda das concessionárias\n"
        "• Time de Inteligência Competitiva Ford BR\n"
        "• Diretoria comercial Ford BR (visão consolidada)\n"
        "• (Indiretamente) o cliente final, que recebe ações "
        "personalizadas mais relevantes"
    ),
    "Estrutura de Custos": (
        "DESENVOLVIMENTO (one-off):\n"
        "• Squad 5 devs × 1 sprint\n\n"
        "OPERAÇÃO MENSAL:\n"
        "• Supabase: ~US$25 (Pro tier)\n"
        "• Hosting API+Web: ~US$20 (Fly/Vercel)\n"
        "• OpenAI/Anthropic: ~US$30-100 (uso por loja)\n"
        "• RapidAPI 411: US$0-19 (free → Pro)\n"
        "• FIPE token: gratuito (1k req/dia)\n\n"
        "Suporte + manutenção: ~10h/mês"
    ),
    "Fontes de Receita": (
        "Para a Ford (redução de custo + aumento de receita):\n"
        "• ↑ VIN Share — clientes retidos geram +R$ 2-5k/ano em pós-venda por veículo\n"
        "• ↓ Custo de aquisição — comparativos rápidos reduzem ciclo de venda\n"
        "• ↑ Conversão de leads frios — ações certas no momento certo\n"
        "• ↓ Custo de pesquisa de mercado manual — automação economiza horas/mês\n\n"
        "Modelo interno (não comercializado externamente)"
    ),
}


def build_canvas_docx():
    doc = Document()

    # Margens compactas pra caber tudo
    for section in doc.sections:
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.orientation = 1  # landscape
        section.page_width, section.page_height = section.page_height, section.page_width

    # Título
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Business Model Canvas — " + PROJETO)
    run.font.size = Pt(16); run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x34, 0x78)

    subt = doc.add_paragraph()
    subt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subt.add_run(CHALLENGE + " · Entrega " + ENTREGA)
    run.font.size = Pt(10); run.italic = True
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    eq = doc.add_paragraph()
    eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = eq.add_run("Equipe: " + " · ".join(f"{n} (RM {r})" for n, r in EQUIPE))
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    doc.add_paragraph()  # spacer

    # Canvas em tabela 3x5 emulando layout clássico
    # Linha 1: Parcerias | Atividades | Proposta de Valor (rowspan 2) | Relacionamento | Segmentos (rowspan 2)
    # Linha 2: (Parcerias cont) | Recursos | (Proposta cont) | Canais | (Segmentos cont)
    # Linha 3: Estrutura de Custos (colspan 2-3) | Fontes de Receita (colspan 2-3)

    # Simplificação prática: tabela 5x5 sem merges complexos — cada célula tem título + conteúdo
    layout = [
        # row 0
        [("Parcerias-Chave", "PARCERIAS-CHAVE", "#003478"),
         ("Atividades-Chave", "ATIVIDADES-CHAVE", "#0066B2"),
         ("Proposta de Valor", "PROPOSTA DE VALOR", "#003478"),
         ("Relacionamento com Clientes", "RELACIONAMENTO COM CLIENTES", "#0066B2"),
         ("Segmentos de Clientes", "SEGMENTOS DE CLIENTES", "#003478")],
        # row 1
        [("", "", ""),  # blank (parcerias continua)
         ("Recursos-Chave", "RECURSOS-CHAVE", "#0066B2"),
         ("", "", ""),  # proposta continua
         ("Canais", "CANAIS", "#0066B2"),
         ("", "", "")],
        # row 2: bottom (custos + receita)
        [("Estrutura de Custos", "ESTRUTURA DE CUSTOS", "#003478"),
         ("Estrutura de Custos", "", "#003478"),  # span continuation
         ("Estrutura de Custos", "", "#003478"),
         ("Fontes de Receita", "FONTES DE RECEITA", "#0066B2"),
         ("Fontes de Receita", "", "#0066B2")],
    ]

    table = doc.add_table(rows=3, cols=5)
    table.style = "Table Grid"
    table.autofit = False

    # Largura uniforme das colunas
    col_widths = [Cm(5.4)] * 5
    for i, w in enumerate(col_widths):
        for cell in table.columns[i].cells:
            cell.width = w

    # Altura linha (em pontos via Tr)
    def set_row_height(row, height_cm):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(int(height_cm * 567)))
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)

    set_row_height(table.rows[0], 7)
    set_row_height(table.rows[1], 5.5)
    set_row_height(table.rows[2], 5)

    # Renderiza
    def fill_cell(cell, key, title_text, color_hex):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        # remove default paragraph
        cell.text = ""
        if title_text:
            p = cell.paragraphs[0]
            run = p.add_run(title_text)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor.from_string(color_hex.lstrip("#"))
            content = CANVAS.get(key, "")
            if content:
                p2 = cell.add_paragraph()
                run2 = p2.add_run(content)
                run2.font.size = Pt(8.5)

    # Row 0
    for col, (key, title, color) in enumerate(layout[0]):
        fill_cell(table.rows[0].cells[col], key, title, color)
    # Row 1
    for col, (key, title, color) in enumerate(layout[1]):
        if title or key not in CANVAS:
            fill_cell(table.rows[1].cells[col], key, title, color)
    # Row 2 — bottom: 3 cells para Custos + 2 cells para Receita
    fill_cell(table.rows[2].cells[0], "Estrutura de Custos", "ESTRUTURA DE CUSTOS", "#003478")
    fill_cell(table.rows[2].cells[3], "Fontes de Receita", "FONTES DE RECEITA", "#0066B2")

    # Merge células visuais: tornar as 3 primeiras da linha 2 uma só, e as 2 últimas outra
    table.rows[2].cells[0].merge(table.rows[2].cells[1]).merge(table.rows[2].cells[2])
    table.rows[2].cells[3].merge(table.rows[2].cells[4])

    # Footer
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = foot.add_run("Documento gerado para a entrega da Sprint Única — Ford × FIAP Challenge 2026.")
    run.font.size = Pt(8); run.italic = True
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    out = OUT / "Business_Canvas.docx"
    doc.save(str(out))
    print(f"OK: {out.name}")


# ============================================================
# 3. QUADRO DE VALOR (.docx)
# ============================================================

QV_ROWS = [
    # Stakeholder, Expectativa, Promessa da solução, Métrica de negócio, Métrica de qualidade, Prioridade
    (
        "Vendedor / Consultor da concessionária",
        "Argumento de venda forte, com dados batidos e comparativos imediatos.",
        "Acesso < 3 cliques ao comparativo de até 5 veículos com fonte verificável de cada spec.",
        "↑ Taxa de conversão de leads (+10 a 20%).",
        "Latência da busca < 3 s, disponibilidade ≥ 99%.",
        "ALTA",
    ),
    (
        "Gerente de Pós-Venda",
        "Saber quais clientes vão evadir antes que evadam, e o que oferecer.",
        "Painel com risco de evasão por cliente + ações sugeridas (4 perfis).",
        "↑ VIN Share da rede oficial (+5 a 8 pp no ano).",
        "Precisão do classificador ≥ 60% (F1 macro); refresh diário.",
        "ALTA",
    ),
    (
        "Time de Inteligência Competitiva Ford BR",
        "Catálogo confiável e auditável dos veículos concorrentes para análise estratégica.",
        "Catálogo com 5+ fontes (FIPE, scraping fabricante, e-book PDF, 411, IA), provenance por campo.",
        "↓ Tempo de pesquisa manual (-70%); base sempre atualizada.",
        "Cobertura ≥ 80% dos modelos BR vigentes; rastreabilidade 100%.",
        "ALTA",
    ),
    (
        "Diretoria Comercial Ford BR",
        "Visão consolidada de saúde da rede e oportunidades por região.",
        "KPIs agregados (vendas, retenção, perfil dominante) por região/dealership.",
        "Decisão estratégica em ≤ 2 dias (vs. semanas de relatórios manuais).",
        "Atualização do dashboard < 24 h; disponibilidade ≥ 99,5%.",
        "MÉDIA",
    ),
    (
        "Cliente final (indireto)",
        "Ofertas e comunicações relevantes para o seu perfil real.",
        "Ações de fidelização (fidelidade premium, lembretes, descontos progressivos) personalizadas.",
        "↑ NPS pós-venda; ↓ churn p/ rede paralela.",
        "Tempo até primeira ação após compra < 7 dias.",
        "ALTA",
    ),
    (
        "Time de Segurança / Compliance Ford",
        "Garantia de proteção de dados pessoais (LGPD) e auditabilidade.",
        "Pseudonimização no pipeline de ML, HMAC nas chamadas, audit log de ações críticas, RLS por dealership.",
        "Zero incidentes de vazamento; LGPD compliance contínuo.",
        "100% requisições ML com payload anonimizado; 100% ações sensíveis auditadas.",
        "ALTA",
    ),
    (
        "Equipe de TI da concessionária",
        "Integração simples com sistemas existentes (CRM, DMS).",
        "API REST documentada (Swagger /docs), SDK opcional, webhooks para eventos.",
        "Tempo de integração < 2 sprints técnicas.",
        "OpenAPI 3.0 completo; uptime ≥ 99%.",
        "MÉDIA",
    ),
]


def build_quadro_valor_docx():
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.orientation = 1
        section.page_width, section.page_height = section.page_height, section.page_width

    # Título
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("Quadro de Valor — " + PROJETO)
    run.bold = True; run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x00, 0x34, 0x78)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = s.add_run(CHALLENGE + " · Entrega " + ENTREGA)
    run.font.size = Pt(10); run.italic = True
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    eq = doc.add_paragraph()
    eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = eq.add_run("Equipe: " + " · ".join(f"{n} (RM {r})" for n, r in EQUIPE))
    run.font.size = Pt(9); run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    doc.add_paragraph()

    headers = ["Stakeholder", "Expectativa", "Promessa da Solução",
               "Métrica de Negócio", "Métrica de Qualidade", "Prioridade"]
    table = doc.add_table(rows=1 + len(QV_ROWS), cols=len(headers))
    table.style = "Table Grid"

    # cabeçalho
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True; run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, "003478")

    # dados
    for r, row in enumerate(QV_ROWS, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            if c == 5:  # prioridade
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.bold = True
                if val == "ALTA":
                    set_cell_bg(cell, "FFE5E5")
                    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
                elif val == "MÉDIA":
                    set_cell_bg(cell, "FFF4D6")
                    run.font.color.rgb = RGBColor(0x8A, 0x65, 0x00)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # ajusta col widths
    widths_cm = [4.5, 5.0, 6.5, 4.0, 5.0, 2.0]
    for i, w in enumerate(widths_cm):
        for cell in table.columns[i].cells:
            cell.width = Cm(w)

    # nota
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "Para cada benefício listado, há uma métrica de NEGÓCIO (impacto monetizável ou "
        "operacional) e uma métrica de QUALIDADE (performance/confiabilidade/SLA), conforme "
        "exigência da disciplina de Testing, Compliance and Quality Assurance."
    )
    run.italic = True; run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    out = OUT / "Quadro_de_Valor.docx"
    doc.save(str(out))
    print(f"OK: {out.name}")


# ============================================================
# 4. README dos entregáveis (.docx)
# ============================================================

def build_readme_docx():
    doc = Document()
    add_heading(doc, "Entregáveis — " + PROJETO, level=0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(CHALLENGE + " · Entrega " + ENTREGA)
    run.italic = True; run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    add_heading(doc, "Equipe", level=1)
    for nome, rm in EQUIPE:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{nome} — RM {rm}").font.size = Pt(11)

    add_heading(doc, "Lista de entregáveis", level=1)
    items = [
        ("Relatorio_Desafio_2_ML.pdf",
         "Relatório técnico do classificador de retenção (Desafio 2): hipótese, bases, pipeline, métricas, perfis e ações."),
        ("Business_Canvas.docx",
         "Business Model Canvas do projeto Faro AI, cobrindo proposta de valor, parcerias, recursos, canais e custos."),
        ("Quadro_de_Valor.docx",
         "Mapeamento de stakeholders × expectativas × métricas de negócio + qualidade × prioridade."),
        ("FaroAI_Architecture.archimate",
         "Modelo de arquitetura no padrão TOGAF/ArchiMate 3.1 — abrir no Archi (https://www.archimatetool.com)."),
        ("../SECURITY.md",
         "Política de segurança cobrindo os 5 eixos avaliativos de Cybersecurity (validação, auth, APIs, dados, monitoramento)."),
    ]
    table = doc.add_table(rows=1 + len(items), cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = ""; hdr[1].text = ""
    for i, name in enumerate(["Arquivo", "Conteúdo"]):
        run = hdr[i].paragraphs[0].add_run(name)
        run.bold = True; run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[i], "003478")
    for r, (name, desc) in enumerate(items, 1):
        c = table.rows[r].cells
        c[0].text = name
        c[1].text = desc
        c[0].paragraphs[0].runs[0].font.size = Pt(10)
        c[0].paragraphs[0].runs[0].font.name = "Consolas"
        c[1].paragraphs[0].runs[0].font.size = Pt(10)
    for cell in table.columns[0].cells:
        cell.width = Cm(7)
    for cell in table.columns[1].cells:
        cell.width = Cm(10)

    add_heading(doc, "Repositório", level=1)
    doc.add_paragraph(
        "Código-fonte completo disponível no monorepo do projeto. "
        "Estrutura principal:"
    )
    for it in [
        "apps/api/ — API gateway Fastify + TS (Desafio 1 + integração ML)",
        "apps/web/ — Front Next.js (admin, busca, comparação, leads)",
        "apps/mobile/ — App Expo (iOS + Android)",
        "services/ml/ — FastAPI Python (classificador + clustering)",
        "supabase/migrations/ — schemas + RLS + audit_log",
        "docs/ — SETUP.md, SECURITY.md, deliverables/",
    ]:
        doc.add_paragraph(it, style="List Bullet")

    add_heading(doc, "Notas de avaliação por disciplina", level=1)
    notes = [
        ("Arquitetura SOA & Web Services", "API REST com Swagger em /docs, separação routes/lib/db, 8 migrations Supabase ordenadas, serviços independentes."),
        ("Mobile Development & IoT", "App Expo 52 + RN 0.76 com expo-router, AsyncStorage, consumo async."),
        ("Testing, Compliance & QA", "Veja Business_Canvas.docx, Quadro_de_Valor.docx, FaroAI_Architecture.archimate. Métricas de negócio + qualidade em todas promessas."),
        ("Cybersecurity", "Veja SECURITY.md — cobre os 5 eixos: validação (Zod), auth (JWT + RBAC), APIs (helmet, rate-limit, CORS, HMAC), dados (pseudonimização, RLS, retenção), monitoramento (Pino + audit_log)."),
        ("IA & Machine Learning", "Veja Relatorio_Desafio_2_ML.pdf + services/ml/notebooks/ford_segmentation.ipynb. Clustering KMeans k=4 + classificador XGBoost sem data leakage."),
    ]
    for d, txt in notes:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(d + ": ")
        run.bold = True; run.font.size = Pt(10.5)
        p.add_run(txt).font.size = Pt(10.5)

    out = OUT / "README_Entregaveis.docx"
    doc.save(str(out))
    print(f"OK: {out.name}")


if __name__ == "__main__":
    import sys
    target = sys.argv[1] if len(sys.argv) > 1 else "all"
    if target in ("all", "pdf"):
        build_ml_report()
    if target in ("all", "canvas"):
        build_canvas_docx()
    if target in ("all", "quadro"):
        build_quadro_valor_docx()
    if target in ("all", "readme"):
        build_readme_docx()
    print(f"\nSalvo em: {OUT}")
